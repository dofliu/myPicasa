
import os
import sys
import markdown
import requests
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class MarkdownToDocxConverter:
    """
    Converts Markdown content to a Docx file with enhanced support for images, tables, and blockquotes.
    """
    def __init__(self):
        self.md = markdown.Markdown(extensions=['markdown.extensions.tables', 'markdown.extensions.fenced_code'])

    def convert(self, md_content, output_path):
        """
        Convert markdown string to docx file.
        :param md_content: Content of the markdown file
        :param output_path: Path to save the docx file
        """
        html = self.md.convert(md_content)
        soup = BeautifulSoup(html, 'html.parser')

        document = Document()
        
        # Iterate over top-level elements
        for element in soup.contents:
            if element.name == 'h1':
                document.add_heading(element.get_text(), level=1)
            elif element.name == 'h2':
                document.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                document.add_heading(element.get_text(), level=3)
            elif element.name == 'h4':
                document.add_heading(element.get_text(), level=4)
            elif element.name == 'h5':
                document.add_heading(element.get_text(), level=5)
            elif element.name == 'h6':
                document.add_heading(element.get_text(), level=6)
            elif element.name == 'p':
                self._add_paragraph_with_images(document, element)
            elif element.name == 'ul':
                self._add_list(document, element, 'List Bullet')
            elif element.name == 'ol':
                self._add_list(document, element, 'List Number')
            elif element.name == 'blockquote':
                self._add_blockquote(document, element)
            elif element.name == 'table':
                self._add_table(document, element)
            elif element.name == 'pre':
                 # Code block usually wrapped in <pre><code>...</code></pre>
                code_text = element.get_text()
                p = document.add_paragraph(code_text)
                p.style = 'No Spacing' 
                # Add a border or shading if desired, but simple monospace is often enough
                for run in p.runs:
                    run.font.name = 'Courier New'

        document.save(output_path)
        return output_path

    def _add_paragraph_with_images(self, document, element):
        """
        Handle paragraphs that might contain images.
        """
        # If no images, just add text
        if not element.find('img'):
             document.add_paragraph(element.get_text())
             return

        # Complex paragraph with images
        # We process children sequentially to maintain order of text and images
        for child in element.children:
            if child.name == 'img':
                img_url = child['src']
                alt_text = child.get('alt', '')
                self._insert_image(document, img_url, alt_text)
            elif child.name is None:
                text = str(child)
                if text.strip():
                    document.add_paragraph(text)
            else:
                 # Other tags like <strong>, <em>, etc.
                 # Python-docx doesn't easily support mixed rich text in one go without building runs manually.
                 # For now, we extract text.
                 text = child.get_text()
                 if text.strip():
                     document.add_paragraph(text)

    def _insert_image(self, document, url, caption):
        """
        Download and insert image.
        """
        try:
            print(f"Downloading image: {url}")
            response = requests.get(url, timeout=10)
            if response.status_code == 200:
                image_stream = BytesIO(response.content)
                document.add_picture(image_stream, width=Inches(5.5)) # Fit to page width approx
                
                if caption:
                    # Add caption
                    caption_p = document.add_paragraph(caption)
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_p.style = 'Caption'
        except Exception as e:
            print(f"Failed to download or insert image {url}: {e}")
            document.add_paragraph(f"[Image: {url} (Download Failed)]")

    def _add_list(self, document, element, style):
        for li in element.find_all('li', recursive=False):
            # recursive=False to avoid nested li items being added to top level
            text = ""
            # Simple text extraction for list items
            # If list item has nested lists, this might need recursion, 
            # but for now let's just get direct text.
            for content in li.contents:
                if content.name in ['ul', 'ol']:
                    # Nested list
                    # Python-docx doesn't strictly support nested list hierarchy easily via styles alone
                    # but we can try just adding them.
                    # Ideally, handle indentation.
                    pass 
                else:
                    text += content.get_text() if content.name else str(content)
            
            clean_text = text.strip()
            if clean_text:
                document.add_paragraph(clean_text, style=style)
            
            # recursive call for nested lists
            for child in li.find_all(['ul', 'ol'], recursive=False):
                 # For nested lists, we might want to change style or just indentation
                 # Word handles indentation by list level, but python-docx needs explicit specific styles
                 # or direct formatting.
                 # Fallback: Just add them as same style for now.
                 self._add_list(document, child, style)

    def _add_blockquote(self, document, element):
        """
        Simulate blockquote with a single-cell table with background.
        """
        text = element.get_text().strip()
        if not text:
            return

        table = document.add_table(rows=1, cols=1)
        # table.style = 'Table Grid' # Optional
        
        cell = table.cell(0, 0)
        cell.text = text
        
        # Add gray shading
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E8E8E8') # Light gray
        tcPr.append(shd)
        
        # Add left border (more complex in oxml)
        # For simplicity, we stick to background color which is the main visual cue.

    def _add_table(self, document, element):
        """
        Convert HTML table to Docx table.
        """
        rows = element.find_all('tr')
        if not rows:
            return
        
        # Determine columns
        first_row = rows[0]
        cols = first_row.find_all(['td', 'th'])
        num_cols = len(cols)
        
        # Create table
        table = document.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < num_cols:
                    table.cell(i, j).text = cell.get_text().strip()

    def convert_file(self, md_file_path, docx_output_path):
        """
        Convert a markdown file to a docx file.
        """
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        return self.convert(md_content, docx_output_path)

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python md2docx_converter.py <input_md_file> [output_docx_file]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        output_file = os.path.splitext(input_file)[0] + ".docx"
    
    converter = MarkdownToDocxConverter()
    try:
        converter.convert_file(input_file, output_file)
        print(f"Successfully converted '{input_file}' to '{output_file}'")
    except Exception as e:
        print(f"Error converting file: {e}")
